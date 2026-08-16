#!/usr/bin/env python3
"""Create a math-doc DOCX skeleton with proof/notes/derivation templates."""
from __future__ import annotations

import argparse
import copy
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from latex_to_omml import latex_to_omml_fixed_alt

BODY_CN = '宋体'
HEADING_CN = '黑体'
TITLE_CN = '方正小标宋简体'
LATIN = 'Times New Roman'
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x59, 0x59, 0x59)
BLACK = RGBColor(0x00, 0x00, 0x00)

TEMPLATES = {
    'proof': {
        'default_title': '证明',
        'sections': [
            ('定理或命题', '在此写出定理、命题或待证明结论。'),
            ('已知条件', '列出已知条件、定义、假设和引用结果。'),
            ('证明', '按逻辑链写出证明过程；关键公式使用编号公式。'),
            ('结论', '回扣待证结论，说明证明完成。'),
        ],
        'citation': '引用格式：[序号]作者.题名.来源.年份.',
        'equation': r'E=mc^{2}',
        'equation_alt': 'E equals m c squared',
    },
    'notes': {
        'default_title': '数学笔记',
        'sections': [
            ('主题', '记录本节主题和复习目标。'),
            ('要点', '用短句列出核心概念、易错点和判断规则。'),
            ('公式', '集中存放本主题公式，并为公式编号。'),
            ('待复习', '列出尚未掌握的题目或知识点。'),
            ('来源', '记录教材、讲义或网页来源。'),
        ],
        'citation': '来源格式：[序号]作者/机构.题名.年份.链接或页码.',
        'equation': r'\int_{-\infty}^{+\infty}\left|x(t)\right|^{2}\,dt',
        'equation_alt': 'E infinity equals integral from minus infinity to plus infinity of absolute value of x of t squared dt',
    },
    'derivation': {
        'default_title': '推导',
        'sections': [
            ('目标', '写出推导目标和最终要得到的形式。'),
            ('假设', '列出模型假设、符号约定和适用范围。'),
            ('推导', '分步推导，每一步注明依据；公式按顺序编号。'),
            ('结果', '汇总最终公式并检查量纲、边界条件和退化情形。'),
            ('引用', '列出推导所依赖的定理、教材或文献。'),
        ],
        'citation': '引用格式：[序号]作者.题名.来源.年份.',
        'equation': r'u(t)=\int_{-\infty}^{t}\delta(\tau)\,d\tau',
        'equation_alt': 'u of t equals integral from minus infinity to t of delta of tau d tau',
    },
}


def set_run_font(run, size=None, bold=None, color=None, east=BODY_CN, latin=LATIN):
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), latin)
    rFonts.set(qn('w:hAnsi'), latin)
    rFonts.set(qn('w:eastAsia'), east)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    sizes = {1: 16, 2: 13, 3: 12}
    colors = {1: BLUE, 2: BLUE, 3: DARK}
    for run in p.runs:
        set_run_font(run, size=sizes.get(level, 12), bold=True, color=colors.get(level, DARK), east=HEADING_CN)
    return p


def add_body(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    set_run_font(run, size=size, color=BLACK)
    return p


def add_numbered_equation(doc, latex, alttext=None, counter=None):
    if counter is None:
        counter = [0]
    counter[0] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    omml = latex_to_omml_fixed_alt(latex, alttext)
    p._element.append(copy.deepcopy(omml))
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run('\t' + f'({counter[0]})')
    set_run_font(run, size=11, color=BLACK)
    return p


def build_docx(template_name: str, title: str, author: str, doc_date: str) -> Document:
    cfg = TEMPLATES[template_name]
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    normal = doc.styles['Normal']
    normal.font.name = LATIN
    normal.font.size = Pt(11)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), LATIN)
    rFonts.set(qn('w:hAnsi'), LATIN)
    rFonts.set(qn('w:eastAsia'), BODY_CN)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    for name, size, color, before, after in (
        ('Heading 1', 16, BLUE, 18, 10),
        ('Heading 2', 13, BLUE, 14, 7),
        ('Heading 3', 12, DARK, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = LATIN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style_rpr = style.element.get_or_add_rPr()
        style_rfonts = style_rpr.get_or_add_rFonts()
        style_rfonts.set(qn('w:eastAsia'), HEADING_CN)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    run = title_p.add_run(title)
    set_run_font(run, size=22, bold=True, color=BLACK)
    title_rpr = run._element.get_or_add_rPr()
    title_rfonts = title_rpr.get_or_add_rFonts()
    title_rfonts.set(qn('w:eastAsia'), TITLE_CN)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(8)
    meta_run = meta.add_run(f'模板：{template_name}；作者：{author}；日期：{doc_date}')
    set_run_font(meta_run, size=10, color=GRAY)

    add_heading(doc, '编号规则', 1)
    add_body(doc, '公式按出现顺序编号：(1)、(2)、(3)...；章节标题使用 Heading 1/2/3，中文标题为黑体，正文为宋体。')
    add_heading(doc, '引用格式', 1)
    add_body(doc, cfg['citation'])

    for heading, body in cfg['sections']:
        add_heading(doc, heading, 1)
        add_body(doc, body)

    add_heading(doc, '示例公式', 1)
    add_numbered_equation(doc, cfg['equation'], alttext=cfg['equation_alt'])
    return doc


def main() -> None:
    parser = argparse.ArgumentParser(description='Create a math-doc DOCX skeleton.')
    parser.add_argument('--template', choices=list(TEMPLATES.keys()), default='notes', help='proof, notes or derivation')
    parser.add_argument('--title', help='document title; defaults to the template title')
    parser.add_argument('--author', default='')
    parser.add_argument('--date', default=str(date.today()))
    parser.add_argument('--output', help='output .docx path')
    args = parser.parse_args()

    cfg = TEMPLATES[args.template]
    title = args.title or cfg['default_title']
    output = Path(args.output or f'mathdoc_{args.template}.docx')
    doc = build_docx(args.template, title, args.author, args.date)
    doc.save(str(output))
    print(output.resolve())


if __name__ == '__main__':
    main()
