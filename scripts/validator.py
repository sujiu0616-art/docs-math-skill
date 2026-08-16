#!/usr/bin/env python3
"""Validate a generated math-doc .docx file.

Usage:
    python validator.py path.docx --level 2

Exit code 0 means the document passed the requested checks.
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

def check_basic(doc: Document) -> list[str]:
    body = doc.element.body
    checks = []

    math_count = sum(1 for _ in body.iter(qn('m:oMath')))
    if math_count == 0:
        raise AssertionError('no OMML equations found')
    checks.append(f'equations={math_count}')

    empty_e = [e for e in body.iter(qn('m:e')) if not e.getchildren()]
    if empty_e:
        raise AssertionError(
            f'{len(empty_e)} empty m:e elements; '
            'check bare |...| absolute values and use \\left|...\\right|'
        )
    checks.append('no empty m:e')

    for ssub in body.iter(qn('m:sSub')):
        if ''.join(ssub.itertext()).startswith('lim'):
            raise AssertionError('lim must use m:limLow, not m:sSub')

    for t in (t.text or '' for t in body.iter(qn('w:t'))):
        if '**' in t:
            raise AssertionError(f"stray ** in text: {t[:60]}")
        if '$' in t:
            has_cjk = any('\u4e00' <= c <= '\u9fff' for c in t)
            if has_cjk and '...' not in t:
                raise AssertionError(f"suspicious $ in CJK context: {t[:60]}")
        if t != t.strip():
            raise AssertionError(f"leading/trailing space in text: {t[:60]}")

    checks.append('no markdown residue')
    checks.append('no leading/trailing spaces')
    return checks


def check_math_layout(doc: Document) -> list[str]:
    body = doc.element.body
    for nary_pr in body.iter(qn('m:naryPr')):
        chr_el = nary_pr.find(qn('m:chr'))
        lim_el = nary_pr.find(qn('m:limLoc'))
        op = chr_el.get(qn('m:val')) if chr_el is not None else None
        expected = 'subSup' if op == '∫' else 'undOvr' if op in ('∑', '∏') else None
        if expected and (lim_el is None or lim_el.get(qn('m:val')) != expected):
            raise AssertionError(f'{op} must use limLoc={expected}')
    return ['nary limits ok']


def check_academic(doc: Document) -> list[str]:
    checks = check_math_layout(doc)
    try:
        normal = doc.styles['Normal']
        rpr = normal.element.find(qn('w:rPr'))
        rfonts = rpr.find(qn('w:rFonts')) if rpr is not None else None
    except KeyError:
        rfonts = None
    east_ok = False
    if rfonts is not None:
        east = rfonts.get(qn('w:eastAsia'))
        east_ok = east == '宋体'
    if not east_ok:
        raise AssertionError('Normal eastAsia must be 宋体')
    checks.append('Normal font ok')

    for name, expected_cn in (('Heading 1', '黑体'), ('Heading 2', '黑体'), ('Heading 3', '黑体')):
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        rpr = style.element.find(qn('w:rPr'))
        rfonts = rpr.find(qn('w:rFonts')) if rpr is not None else None
        if rpr is None or rfonts is None or rfonts.get(qn('w:eastAsia')) != expected_cn:
            raise AssertionError(f'{name} eastAsia must be {expected_cn}')
        if name == 'Heading 2' and style.font.bold is not True:
            raise AssertionError('Heading 2 must be bold')
    checks.append('Heading fonts ok')

    if 'Title' in doc.styles:
        style = doc.styles['Title']
        rpr = style.element.find(qn('w:rPr'))
        rfonts = rpr.find(qn('w:rFonts')) if rpr is not None else None
        if rpr is None or rfonts is None or rfonts.get(qn('w:eastAsia')) != '方正小标宋简体':
            raise AssertionError('Title eastAsia must be 方正小标宋简体')
    checks.append('Title font ok')

    for table in doc.tables:
        grid = table._tbl.tblGrid
        cols = grid.findall(qn('w:gridCol')) if grid is not None else []
        actual = [int(c.get(qn('w:w'), 0)) for c in cols]
        if len(actual) != len(table.columns):
            raise AssertionError('tblGrid column count mismatch')
        if not all(w > 0 for w in actual):
            raise AssertionError('tblGrid contains zero-width columns')

        rows = table._tbl.findall(qn('w:tr'))
        if not rows:
            continue
        first_tr_pr = rows[0].find(qn('w:trPr'))
        if first_tr_pr is None or first_tr_pr.find(qn('w:tblHeader')) is None:
            raise AssertionError('first table row must repeat as header')
        if first_tr_pr.find(qn('w:cantSplit')) is None:
            raise AssertionError('header row must be cantSplit')

        for tc in rows[0].findall(qn('w:tc')):
            tc_pr = tc.find(qn('w:tcPr'))
            if tc_pr is None:
                continue
            shd = tc_pr.find(qn('w:shd'))
            if shd is None or shd.get(qn('w:fill')) != 'E8EEF5':
                raise AssertionError('header cells must use fill E8EEF5')
            v_align = tc_pr.find(qn('w:vAlign'))
            if v_align is None or v_align.get(qn('w:val')) != 'center':
                raise AssertionError('cells must use vAlign=center')

        for row in rows[1:]:
            for tc in row.findall(qn('w:tc')):
                tc_pr = tc.find(qn('w:tcPr'))
                if tc_pr is None:
                    continue
                v_align = tc_pr.find(qn('w:vAlign'))
                if v_align is None or v_align.get(qn('w:val')) != 'center':
                    raise AssertionError('cells must use vAlign=center')

    checks.append(f'tables={len(doc.tables)} ok')
    return checks


def validate_docx(path: str | Path, level: int = 1) -> list[str]:
    doc_path = Path(path)
    if not doc_path.exists():
        raise FileNotFoundError(doc_path)

    doc = Document(str(doc_path))
    checks = check_basic(doc)
    if level >= 2:
        checks.extend(check_academic(doc))
    if level >= 3:
        for section in doc.sections:
            if section.page_width.cm < 10 or section.page_height.cm < 10:
                raise AssertionError('page size looks invalid')
        checks.append('page size ok')
    return checks


def main() -> None:
    parser = argparse.ArgumentParser(description='Validate math-doc .docx output.')
    parser.add_argument('docx', type=Path, help='generated .docx file')
    parser.add_argument('--level', type=int, choices=(1, 2, 3), default=1)
    args = parser.parse_args()

    try:
        checks = validate_docx(args.docx, level=args.level)
    except (AssertionError, FileNotFoundError) as exc:
        print(f'FAIL: {exc}')
        raise SystemExit(1)
    print('OK: ' + ', '.join(checks))


if __name__ == '__main__':
    main()
