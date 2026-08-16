# -*- coding: utf-8 -*-
"""Reusable OMML builders aligned with docx-math SKILL.md API.

Usage:
    from omml_helpers import *
    doc = Document()
    set_normal_style(doc)
    mpara(doc, [mrn("S="), mnary("∑", [mrn("n=0")], [mrn("∞")], [msup([mr("α")], [mr("n")])])])
    doc.save("out.docx")
"""
import copy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

FONT_CN = '宋体'        # SimSun
FONT_CN_FALLBACK = '宋体'        # SimSun fallback
FONT_HEADING_CN = '黑体'          # SimHei
FONT_MATH = 'Times New Roman'


# ===== Element factory =====
def E(tag, text=None):
    """Create an OxmlElement with optional text content."""
    el = OxmlElement(tag)
    if text is not None:
        el.text = text
    return el


def _mv(el, val):
    """Set m:val attribute on an element."""
    el.set(qn("m:val"), val)


# ===== Math runs =====
def mr(text, italic=False):
    """Math run: <m:r><m:t>text</m:t></m:r>, italic by convention."""
    r = E("m:r")
    rPr = E("m:rPr")
    if italic:
        it = E("m:ital"); _mv(it, "1"); rPr.append(it)
    scr = E("m:scr"); _mv(scr, "latin"); rPr.append(scr)
    r.append(rPr)
    t = E("m:t", text); t.set(qn("xml:space"), "preserve"); r.append(t)
    return r


def mrn(text):
    """Non-italic math run (for =, +, numbers, etc.)."""
    r = E("m:r"); rPr = E("m:rPr"); r.append(rPr)
    t = E("m:t", text); t.set(qn("xml:space"), "preserve"); r.append(t)
    return r


def meq():
    """Equals sign, no surrounding spaces (SKILL.md Lesson #6)."""
    return mrn("=")


# ===== Builders (all deepcopy children — safe to reuse) =====
def msup(bl, sl):
    """Superscript: base^sup."""
    ss = E("m:sSup"); e = E("m:e"); su = E("m:sup")
    for x in bl: e.append(copy.deepcopy(x))
    for x in sl: su.append(copy.deepcopy(x))
    ss.append(e); ss.append(su); return ss


def msub(bl, sl):
    """Subscript: base_sub."""
    ss = E("m:sSub"); e = E("m:e"); sub = E("m:sub")
    for x in bl: e.append(copy.deepcopy(x))
    for x in sl: sub.append(copy.deepcopy(x))
    ss.append(e); ss.append(sub); return ss


def msubsup(bl, subl, supl):
    """Sub+Superscript: base_{sub}^{sup}."""
    sss = E("m:sSubSup"); e = E("m:e")
    sub = E("m:sub"); su = E("m:sup")
    for x in bl:   e.append(copy.deepcopy(x))
    for x in subl: sub.append(copy.deepcopy(x))
    for x in supl: su.append(copy.deepcopy(x))
    sss.append(e); sss.append(sub); sss.append(su); return sss


def mfrac(nl, dl):
    """Fraction: num/den."""
    f = E("m:f"); n = E("m:num"); d = E("m:den")
    for x in nl: n.append(copy.deepcopy(x))
    for x in dl: d.append(copy.deepcopy(x))
    f.append(n); f.append(d); return f


def mnary(chr_char, lo_list, hi_list, body_list):
    """N-ary operator (∑ ∏ ∫ …) with limits above/below (limLoc=undOvr)."""
    nary = E("m:nary"); naryPr = E("m:naryPr")
    chr_el = E("m:chr"); _mv(chr_el, chr_char); naryPr.append(chr_el)
    limLoc = E("m:limLoc"); _mv(limLoc, "undOvr"); naryPr.append(limLoc)
    nary.append(naryPr)
    sub_el = E("m:sub")
    for x in lo_list: sub_el.append(copy.deepcopy(x))
    nary.append(sub_el)
    sup_el = E("m:sup")
    for x in hi_list: sup_el.append(copy.deepcopy(x))
    nary.append(sup_el)
    e_el = E("m:e")
    for x in body_list: e_el.append(copy.deepcopy(x))
    nary.append(e_el)
    return nary


def mlim(bl, ll):
    """Lower limit: base with limit below (e.g. lim_{m→∞})."""
    lm = E("m:limLow"); e = E("m:e"); lim = E("m:lim")
    for x in bl: e.append(copy.deepcopy(x))
    for x in ll: lim.append(copy.deepcopy(x))
    lm.append(e); lm.append(lim); return lm


def mdel(content_lists, beg='(', end=')'):
    """Delimited group: (content). beg/end configurable for [] {} etc."""
    d = E("m:d"); dPr = E("m:dPr")
    b = E("m:begChr"); _mv(b, beg); dPr.append(b)
    e_chr = E("m:endChr"); _mv(e_chr, end); dPr.append(e_chr)
    d.append(dPr)
    for ne in content_lists:
        e_el = E("m:e")
        lst = ne if isinstance(ne, list) else [ne]
        for x in lst: e_el.append(copy.deepcopy(x))
        d.append(e_el)
    return d


def mabs(content_lists):
    """Absolute value: |content|."""
    return mdel(content_lists, beg='|', end='|')


def omath(*c):
    """Wrap children in <m:oMath>."""
    o = E("m:oMath")
    for x in c: o.append(copy.deepcopy(x))
    return o


# ===== Chinese font helper =====
def _set_cn(run, font=FONT_CN):
    """Set East Asian font on a run (SKILL.md Core Lesson #5)."""
    rp = run._element.get_or_add_rPr()
    rf = rp.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rp.insert(0, rf)
    rf.set(qn('w:eastAsia'), font)


def set_normal_style(doc, font_name=FONT_MATH, font_size=Pt(12),
                     eastasia=FONT_CN, line_spacing=1.5, space_after=Pt(6)):
    """Configure Normal style: TNR + SimSun, so math and Chinese coexist."""
    sn = doc.styles['Normal']
    sn.font.name = font_name
    sn.font.size = font_size
    sn.paragraph_format.line_spacing = line_spacing
    sn.paragraph_format.space_after = space_after
    rPr = sn.element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr'); sn.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), eastasia)
    return sn


# ===== Paragraph builders (doc passed in — pure module) =====
def mpara(doc, body, align=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(8)):
    """Pure-math paragraph. CRITICAL: appends oMath to p._element (SKILL.md Lesson #7)."""
    p = doc.add_paragraph(); p.alignment = align
    pPr = p._element.get_or_add_pPr()
    sp = E("w:spacing")
    sp.set(qn("w:before"), "0")
    sp.set(qn("w:after"), str(int(sa / Pt(1) * 20)))
    pPr.append(sp)
    p._element.append(omath(*body))
    return p


def mpara_mix(doc, segments, align=WD_ALIGN_PARAGRAPH.LEFT,
              sa=Pt(6), indent=False):
    """Mixed text + math paragraph.
    segments: list of (kind, data)
        ('t', text)   — normal text run
        ('b', text)   — bold text run
        ('m', [omath children]) — inline math
    """
    p = doc.add_paragraph(); p.alignment = align
    pPr = p._element.get_or_add_pPr()
    sp = E("w:spacing")
    sp.set(qn("w:before"), "0")
    sp.set(qn("w:after"), str(int(sa / Pt(1) * 20)))
    pPr.append(sp)
    if indent:
        p.paragraph_format.left_indent = Inches(0.4)
    for kind, data in segments:
        if kind == 't':
            r = p.add_run(data)
            r.font.name = FONT_MATH; r.font.size = Pt(12)
            _set_cn(r)
        elif kind == 'b':
            r = p.add_run(data)
            r.font.name = FONT_MATH; r.font.size = Pt(12)
            r.bold = True
            _set_cn(r)
        elif kind == 'm':
            p._element.append(omath(*data))
    return p
