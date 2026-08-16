#!/usr/bin/env python3
"""Generate the four example documents: lecture notes, proof, exercise set, paper.

Run from the repository root:
    python examples/generate_samples.py
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / 'scripts'))

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Inches, Pt  # noqa: E402

from mathdoc_cli import (  # noqa: E402
    BODY_CN, HEADING_CN, LATIN, TITLE_CN, add_body, add_heading,
    add_numbered_equation, build_docx, set_run_font,
)

OUT = Path(__file__).resolve().parent


def make_doc(title: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    normal = doc.styles['Normal']
    normal.font.name = LATIN
    normal.font.size = Pt(11)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', LATIN)
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', LATIN)
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', BODY_CN)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    for name, size, before, after in (
        ('Heading 1', 16, 18, 10),
        ('Heading 2', 13, 14, 7),
        ('Heading 3', 12, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = LATIN
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.element.get_or_add_rPr().get_or_add_rFonts().set(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', HEADING_CN)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    run = title_p.add_run(title)
    set_run_font(run, size=22, bold=True, color=None)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', TITLE_CN)
    return doc


def build_exercises() -> Document:
    doc = make_doc('微积分习题集（示例）')
    add_heading(doc, '题目', 1)
    add_heading(doc, '第一题：定积分', 2)
    add_body(doc, '计算定积分并给出依据：')
    add_numbered_equation(doc, r'\int_{0}^{1}x^{2}\,dx')
    add_heading(doc, '第二题：级数', 2)
    add_body(doc, '判断级数的收敛性并求其和：')
    add_numbered_equation(doc, r'\sum_{n=1}^{\infty}\frac{1}{n(n+1)}')
    add_heading(doc, '第三题：极限', 2)
    add_body(doc, '求极限：')
    add_numbered_equation(doc, r'\lim_{x\to 0}\frac{\sin x}{x}')
    add_heading(doc, '第四题：绝对值不等式', 2)
    add_body(doc, '证明对所有实数 a、b 成立：')
    add_numbered_equation(doc, r'\left|a+b\right|\le\left|a\right|+\left|b\right|')
    add_heading(doc, '参考答案提示', 1)
    add_body(doc, '第一题利用牛顿-莱布尼茨公式；第二题裂项相消；第三题利用重要极限；第四题分符号讨论。')
    return doc


def build_paper() -> Document:
    doc = make_doc('关于序列变换性质的示例论文')
    add_heading(doc, '摘要', 1)
    add_body(doc, '本文以示例方式展示含编号公式的论文结构：介绍序列的变换表示，导出闭合形式并讨论其性质。')
    add_heading(doc, '1 引言', 1)
    add_body(doc, '对离散时间序列，变换是分析系统行为的核心工具。本文给出两个基本结果。')
    add_heading(doc, '2 公式与结果', 1)
    add_heading(doc, '2.1 变换的定义', 2)
    add_body(doc, '序列 x[n] 的变换定义为：')
    add_numbered_equation(doc, r'X(z)=\sum_{n=-\infty}^{+\infty}x[n]z^{-n}')
    add_heading(doc, '2.2 几何级数求和', 2)
    add_body(doc, '当收敛域包含单位圆时，几何级数可化为闭合形式：')
    add_numbered_equation(doc, r'\sum_{n=0}^{+\infty}\alpha^{n}=\frac{1}{1-\alpha},\qquad\left|\alpha\right|<1')
    add_heading(doc, '2.3 逆变换', 2)
    add_body(doc, '逆变换由围线积分给出：')
    add_numbered_equation(doc, r'x[n]=\frac{1}{2\pi j}\oint X(z)z^{n-1}\,dz')
    add_heading(doc, '3 结论', 1)
    add_body(doc, '变换与逆变换互为反演，公式 (1)-(3) 提供完整的计算路径。')
    return doc


def main() -> None:
    (build_docx('notes', '信号与系统讲义（示例）', 'math-doc', '')
     .save(str(OUT / 'lecture_notes.docx')))
    (build_docx('proof', '勾股定理的证明（示例）', 'math-doc', '')
     .save(str(OUT / 'proof.docx')))
    build_exercises().save(str(OUT / 'exercises.docx'))
    build_paper().save(str(OUT / 'paper.docx'))
    print('generated: lecture_notes.docx, proof.docx, exercises.docx, paper.docx')


if __name__ == '__main__':
    main()
